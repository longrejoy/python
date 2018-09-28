# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'D:\python\GUI\Editor.ui'
#
# Created by: PyQt5 UI code generator 5.11.2
#
# WARNING! All changes made in this file will be lost!

#from PyQt5 import *
import sys

import docx
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from docx import Document
from docx.shared import Inches

import Editor_qrc


class Ui_MainWindow(object):

    def setupUi(self, MainWinow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(320, 240)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.frame = QtWidgets.QFrame(self.centralwidget)
        self.frame.setGeometry(QtCore.QRect(0, 0, 320, 240))
        self.frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        self.textBrowser = QtWidgets.QTextBrowser(self.frame)
        self.textBrowser.setGeometry(QtCore.QRect(0, 0, 320, 240))
        self.textBrowser.setObjectName("textBrowser")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 320, 26))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        self.menu_2 = QtWidgets.QMenu(self.menubar)
        self.menu_2.setObjectName("menu_2")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.action = QtWidgets.QAction(MainWindow)
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(":/menu/open.jpg"), QtGui.QIcon.Normal, QtGui.QIcon.On)
        self.action.setIcon(icon)
        self.action.setObjectName("action")
        self.action_2 = QtWidgets.QAction(MainWindow)
        icon1 = QtGui.QIcon()
        icon1.addPixmap(QtGui.QPixmap(":/menu/save.jpg"), QtGui.QIcon.Normal, QtGui.QIcon.On)
        self.action_2.setIcon(icon1)
        self.action_2.setObjectName("action_2")
        self.action_3 = QtWidgets.QAction(MainWindow)
        icon2 = QtGui.QIcon()
        icon2.addPixmap(QtGui.QPixmap(":/menu/exit.jpg"), QtGui.QIcon.Normal, QtGui.QIcon.On)
        self.action_3.setIcon(icon2)
        self.action_3.setObjectName("action_3")
        self.action_4 = QtWidgets.QAction(MainWindow)
        icon3 = QtGui.QIcon()
        icon3.addPixmap(QtGui.QPixmap(":/menu/about.jpg"), QtGui.QIcon.Normal, QtGui.QIcon.On)
        self.action_4.setIcon(icon3)
        self.action_4.setObjectName("action_4")
        self.menu.addAction(self.action)
        self.menu.addAction(self.action_2)
        self.menu.addAction(self.action_3)
        self.menu_2.addAction(self.action_4)
        self.menubar.addAction(self.menu.menuAction())
        self.menubar.addAction(self.menu_2.menuAction())

        self.retranslateUi(MainWindow)
        self.action.triggered.connect(self.on_action_triggered)
        self.action_2.triggered.connect(self.on_action_2_triggered)
        self.action_3.triggered.connect(self.on_action_3_triggered)
        self.action_4.triggered.connect(self.on_action_4_triggered)
        #self.textBrowser.anchorClicked.connect(self.on_textBrowser_clicked('content'))

        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    #def on_textBrowser_clicked(self,content):
    #    self.textBrowser.textCursor().insertText(self.textBrowser.textCursor().insertText,content,'')

    def on_action_triggered(self):
        print(u"打开")
        my_file_path = QtWidgets.QFileDialog.getOpenFileName(self.centralwidget,u'打开文件','/')
        print((my_file_path[0]))
        print(type(my_file_path[0]))
        my_file_path_1 = my_file_path[0].replace('/','\\')
        print(my_file_path_1)
        print(type(my_file_path_1))
        doc = docx.Document(my_file_path_1)
        self.textBrowser.clear()
        for my_paragraph in doc.paragraphs:
            print(my_paragraph.text)
            self.textBrowser.append(my_paragraph.text)

    def on_action_2_triggered(self):
        print(u"保存")
        my_file_path = QtWidgets.QFileDialog.getSaveFileName(self.centralwidget,u'保存文件','/')
        my_file_path_1 = my_file_path[0].replace('/','\\')
        print(my_file_path_1)
        #doc = docx.Document(my_file_path_1)
        print('11111')
        print(self.textBrowser.toPlainText())
        #p = Document.add_paragraph(self.textBrowser.toPlainText())
        context = docx.Document()
        p = context.add_paragraph(self.textBrowser.toPlainText())
        context.save(my_file_path_1)
        

    def on_action_3_triggered(self):
        print(u"退出")
        exit(0)

    def on_action_4_triggered(self):
        print(u"关于")
        print(type(Ui_MainWindow))
        print(type(MainWindow))
        My_Message = QMessageBox.about(self.centralwidget,u'关于',u'这是一个文本编辑器')
        

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.menu.setTitle(_translate("MainWindow", "文件"))
        self.menu_2.setTitle(_translate("MainWindow", "帮助"))
        self.action.setText(_translate("MainWindow", "打开"))
        self.action_2.setText(_translate("MainWindow", "保存"))
        self.action_3.setText(_translate("MainWindow", "退出"))
        self.action_4.setText(_translate("MainWindow", "关于"))


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
