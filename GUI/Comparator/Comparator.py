from PyQt5 import QtCore, QtGui, QtWidgets
from Ui_Comparator import Ui_MainWindow
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QFileDialog
from PyQt5.QtWidgets import QTableWidget
from PyQt5.QtWidgets import QTableWidgetItem
from PyQt5.QtWidgets import QMessageBox

from docx import Document
import docx
import xlwt
import xlrd
import Qrc_Comparator
class MainWindow(QMainWindow,Ui_MainWindow):
    FileType1 = ''
    FileType2 = ''
    def __init__(self, parent = None):
        QMainWindow.__init__(self, parent)
        self.setupUi(self)

    @pyqtSlot()
    def on_pushButton_clicked(self):
        print("Open File1")
        file_name = QtWidgets.QFileDialog.getOpenFileName(self,"Open File1","./")
        file_path_new = file_name[0].replace('/','\\')
        print(type(file_path_new))
        if file_path_new[-5:] == '.docx':
            self.FileType1 = 'docx'
            doc = docx.Document(file_path_new)
            self.textBrowser_2.setVisible(True)
            self.tableWidget.setVisible(False)
            self.textBrowser_2.clear()
            for my_paragraph in doc.paragraphs:
                print(my_paragraph.text)    
                self.textBrowser_2.append(my_paragraph.text)
        elif file_path_new[-5:] == '.xlsx':
            self.FileType1 = 'xlsx'
            self.tableWidget.clear()
            excel = xlrd.open_workbook(file_path_new)
            self.textBrowser_2.setVisible(False)
            self.tableWidget.setVisible(True)
            #self.tableWidget.horizontalHeader().setResizeMode(QHeaderView.Interactive)
            #self.tableWidget.setAutoScroll(True)
            print(excel.sheet_names())
            table = excel.sheets()[0]
            print(table)
            self.tableWidget.setColumnCount(table.ncols)
            self.tableWidget.setRowCount(table.nrows)
            for i in range(table.nrows):
                for j in range(table.ncols):
                    if table.cell(i,j).value:
                        print(table.cell(i,j).value)
                        newItem = QTableWidgetItem(str(table.cell(i,j).value))
                        #newItem = QTableWidgetItem('table.cell(i,j).value')
                        self.tableWidget.setItem(i,j,newItem)
        elif file_path_new[-4:] == '.txt':
            self.FileType1 = 'txt'
            f = open(file_path_new, 'r')
            self.textBrowser_2.clear()
            self.textBrowser_2.setVisible(True)
            self.tableWidget.setVisible(False)
            with f:
                data = f.read()
                self.textBrowser_2.append(data)
            f.close()

    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        print("Open File1")
        file_name = QtWidgets.QFileDialog.getOpenFileName(self,"Open File2","./")
        file_path_new = file_name[0].replace('/','\\')
        print(type(file_path_new))
        if file_path_new[-5:] == '.docx':
            self.FileType2 = 'docx'
            doc = docx.Document(file_path_new)
            self.textBrowser_3.setVisible(True)
            self.tableWidget_2.setVisible(False)
            self.textBrowser_3.clear()
            for my_paragraph in doc.paragraphs:
                print(my_paragraph.text)    
                self.textBrowser_3.append(my_paragraph.text)
        elif file_path_new[-5:] == '.xlsx':
            self.FileType2 = 'xlsx'
            excel = xlrd.open_workbook(file_path_new)
            self.textBrowser_3.setVisible(False)
            self.tableWidget_2.setVisible(True)
            self.tableWidget_2.clear()
            #self.tableWidget_2.horizontalHeader().setResizeMode(QHeaderView.Interactive)
            #self.tableWidget_2.setAutoScroll(True)
            print(excel.sheet_names())
            table = excel.sheets()[0]
            print(table)
            self.tableWidget_2.setColumnCount(table.ncols)
            self.tableWidget_2.setRowCount(table.nrows)
            for i in range(table.nrows):
                for j in range(table.ncols):
                    if table.cell(i,j).value:
                        print(table.cell(i,j).value)
                        newItem = QTableWidgetItem(str(table.cell(i,j).value))
                        #newItem = QTableWidgetItem('table.cell(i,j).value')
                        self.tableWidget_2.setItem(i,j,newItem)
        elif file_path_new[-4:] == '.txt':
            self.FileType2 = 'txt'
            f = open(file_path_new, 'r')
            self.textBrowser_3.clear()
            self.textBrowser_3.setVisible(True)
            self.tableWidget_2.setVisible(False)
            with f:
                data = f.read()
                self.textBrowser_3.append(data)
            f.close()
    
    @pyqtSlot()
    def on_pushButton_3_clicked(self):
        print("start compare")
        self.textBrowser.clear()
        if self.FileType1 != self.FileType2:
            print('unmatch file type')
        elif self.FileType1 == 'docx':
            print("docx")
            content1 = self.textBrowser_2.toPlainText()
            content2 = self.textBrowser_3.toPlainText()
            print(type(content1))
            str1 = content1.split('\n')
            str2 = content2.split('\n')
            dif_count = 0
            for str_count in range(len(str1)):
                if str1[str_count] != str2[str_count]:
                    self.textBrowser.append("different line:" +  str(str_count + 1))
                    self.textBrowser.append("File1 is:")
                    self.textBrowser.append(str1[str_count])
                    self.textBrowser.append("File2 is:")
                    self.textBrowser.append(str2[str_count])
                    dif_count += 1
            self.textBrowser.append("Total different lines:" + str(dif_count))
        elif self.FileType1 == 'xlsx':
            print('xlsx')
            dif_count = 0
            for i in range(self.tableWidget.rowCount()):
                for j in range(self.tableWidget.columnCount()):
                    #print(self.tableWidget.item(i,j).text())
                    if (self.tableWidget.item(i,j)) and (self.tableWidget_2.item(i,j)):
                        if (self.tableWidget.item(i,j).text() == self.tableWidget_2.item(i,j).text()):
                            print(' ')
                        else:
                            self.textBrowser.append("different row" + str(i+1) + " column" + str(j+1))
                            self.textBrowser.append("File1 is:")
                            self.textBrowser.append(str(self.tableWidget.item(i,j).text()))
                            self.textBrowser.append("File2 is:")
                            self.textBrowser.append(str(self.tableWidget_2.item(i,j).text()))
                            dif_count += 1
                    elif (not self.tableWidget.item(i,j)) and (not self.tableWidget_2.item(i,j)):
                        break
                    else:
                        self.textBrowser.append("different row" + str(i+1) + " column" + str(j+1))
                        self.textBrowser.append("File1 is:")
                        if self.tableWidget.item(i,j):
                            self.textBrowser.append(str(self.tableWidget.item(i,j).text()))
                        else:
                            self.textBrowser.append("null")
                        self.textBrowser.append("File2 is:")
                        if self.tableWidget_2.item(i,j):
                            self.textBrowser.append(str(self.tableWidget_2.item(i,j).text()))
                        else:                        
                            self.textBrowser.append("null")
                        dif_count += 1
            self.textBrowser.append("Total different lines:" + str(dif_count))
        elif self.FileType1 == 'txt':
            print("txt")
            content1 = self.textBrowser_2.toPlainText()
            content2 = self.textBrowser_3.toPlainText()
            print(type(content1))
            str1 = content1.split('\n')
            str2 = content2.split('\n')
            dif_count = 0
            for str_count in range(len(str1)):
                if str1[str_count] != str2[str_count]:
                    self.textBrowser.append("different line:" +  str(str_count + 1))
                    self.textBrowser.append("File1 is:")
                    self.textBrowser.append(str1[str_count])
                    self.textBrowser.append("File2 is:")
                    self.textBrowser.append(str2[str_count])
                    dif_count += 1
            self.textBrowser.append("Total different lines:" + str(dif_count))

    @pyqtSlot()
    def on_action_3_triggered(self):
        print(u"退出")
        sys.exit(app.exec_())

    @pyqtSlot()
    def on_action_4_triggered(self):
        print(u"关于")
        print(type(Ui_MainWindow))
        print(type(MainWindow))
        My_Message = QMessageBox.about(self.centralwidget,u'关于',u'这是一个文本比较器，目前支持txt、docx和xlsx文件比较')

if __name__ == "__main__":
    import sys
    ##app = QtGui.QApplication(sys.argv)
    #error:app = QtGui.QGuiApplication(sys.argv)
    app = QtWidgets.QApplication(sys.argv)
    ##MainWindow = QtGui.QMainWindow()
    #MainWindow = QtWidgets.QMainWindow()
    ui = MainWindow()
    #ui.setupUi(MainWindow)
    ui.show()
    sys.exit(app.exec_())